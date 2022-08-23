from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.config import DOCUMENT_TEXT, DOCUMENT_TEXT_HEAT, DOCUMENT_TEXT_LOAD


class WordDocument:
    def __init__(self,
                 path_file="Test.docx",
                 name_file="Test",
                 name="Test",
                 date_in="01/01/01 01:01",
                 date_out="99/99/99 99:99",
                 type_chart="Линейный",
                 param_draw="Данные, как есть"):
        self.path = path_file
        self.name_file = name_file
        self.name = name

        self.date_in = date_in
        self.date_out = date_out

        self.type_chart = type_chart
        self.param_draw = param_draw

        self.document = Document()

    def set_style(self):
        self.style_text = self.document.styles.add_style("Text_main", WD_STYLE_TYPE.PARAGRAPH)
        self.style_text.font.name = "Arial"
        self.style_text.font.size = Pt(14)

        self.style_head = self.document.styles.add_style("Head_main", WD_STYLE_TYPE.PARAGRAPH)
        self.style_head.font.name = "Arial"
        self.style_head.font.size = Pt(16)
        self.style_head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def set_title(self):
        self.document.add_paragraph(f"Отчёт по прибору: {self.name}", style="Head_main")
        self.document.add_paragraph(f"За промежуток с {self.date_in} по {self.date_out}", style="Head_main")

    def set_main_data(self, image):
        self.main_image = image

        self.document.add_paragraph(DOCUMENT_TEXT.format(
            self.name,
            self.date_in,
            self.date_out,
            self.type_chart,
            self.param_draw
        ), style="Text_main")
        self.document.add_picture(self.main_image, width=Mm(150))

    def set_heat_data(self, image):
        self.heat_image = image

        self.document.add_paragraph(DOCUMENT_TEXT_HEAT.format(
            self.name,
            self.date_in,
            self.date_out,
        ), style="Text_main")
        self.document.add_picture(self.heat_image, width=Mm(150))

    def set_load_data(self, image):
        self.load_data = image

        self.document.add_paragraph(DOCUMENT_TEXT_LOAD.format(
            self.name,
            self.date_in,
            self.date_out,
        ), style="Text_main")
        self.document.add_picture(self.load_data, width=Mm(150))

    def save(self):
        self.document.save(self.path)


if __name__ == '__main__':
    document = WordDocument()

    document.set_style()
    document.set_title()

    # document.set_main_data()
    # document.set_heat_data()
    # document.set_load_data()

    document.save()
